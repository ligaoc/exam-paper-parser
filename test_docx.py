"""快速测试docx文件是否可以直接提取文本"""
from docx import Document

# 读取docx文件
doc_path = r"D:\商业项目\试卷ocr\2022年广东省广州市中考化学真题（原卷版）.docx"
doc = Document(doc_path)

print("=" * 60)
print("文档基本信息")
print("=" * 60)
print(f"段落数量: {len(doc.paragraphs)}")
print(f"表格数量: {len(doc.tables)}")

print("\n" + "=" * 60)
print("前20个段落内容预览")
print("=" * 60)
for i, para in enumerate(doc.paragraphs[:20]):
    text = para.text.strip()
    if text:
        print(f"[{i}] {text[:100]}{'...' if len(text) > 100 else ''}")

print("\n" + "=" * 60)
print("检测目录编号、括号、下划线、分数")
print("=" * 60)
import re

all_text = "\n".join([p.text for p in doc.paragraphs])

# 检测括号
brackets = re.findall(r'[（(][^）)]*[）)]', all_text)
print(f"\n括号内容示例 (前10个): {brackets[:10]}")

# 检测分数
scores = re.findall(r'[（(【\[]?\d+分[）)\]】]?', all_text)
print(f"\n分数示例: {scores[:15]}")

# 检测下划线
underlines = re.findall(r'_{2,}', all_text)
print(f"\n下划线数量: {len(underlines)}")

# 检测可能的题号
numbers = re.findall(r'^[\d一二三四五六七八九十]+[、.．]', all_text, re.MULTILINE)
print(f"\n题号示例: {numbers[:15]}")
